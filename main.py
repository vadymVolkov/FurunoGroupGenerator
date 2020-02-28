import sys  # sys нужен для передачи argv в QApplication
from PyQt5 import QtWidgets
import designe
import datetime
import os
from docx import Document
from docx import shared
import openpyxl
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PyPDF2 import PdfFileWriter, PdfFileReader


class MainApp(QtWidgets.QMainWindow, designe.Ui_MainWindow):
    day1 = datetime.datetime.now()
    day2 = day1 + datetime.timedelta(days=1)
    name_file_path = './Documents/names.docx'
    work_file_path = os.path.abspath(os.getcwd()) + '/Documents/'

    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.exit_button.clicked.connect(self.close)
        day1_str = self.day1.strftime("%d.%m.%Y")
        day2_str = self.day2.strftime("%d.%m.%Y")
        self.day1_value.setText(str(day1_str))
        self.day2_value.setText(str(day2_str))
        self.day_backward_button.clicked.connect(self.day_backward)
        self.day_forward_button.clicked.connect(self.day_forward)
        self.browse_group_folder_button.clicked.connect(self.select_group_folder)
        self.create_new_group_button.clicked.connect(self.generate_group_folder)
        self.path_group_folder_value.setText(os.path.abspath(os.getcwd()) + '/' + self.folder_name())
        self.generate_test_files_button.clicked.connect(self.generate_test_files)

    def day_backward(self):
        self.day1 = self.day1 - datetime.timedelta(days=1)
        self.day2 = self.day1 + datetime.timedelta(days=1)
        day1_str = self.day1.strftime("%d.%m.%Y")
        day2_str = self.day2.strftime("%d.%m.%Y")
        self.day1_value.setText(str(day1_str))
        self.day2_value.setText(str(day2_str))
        self.path_group_folder_value.setText(os.path.abspath(os.getcwd()) + '/' + self.folder_name())

    def day_forward(self):
        self.day1 = self.day1 + datetime.timedelta(days=1)
        self.day2 = self.day1 + datetime.timedelta(days=1)
        day1_str = self.day1.strftime("%d.%m.%Y")
        day2_str = self.day2.strftime("%d.%m.%Y")
        self.day1_value.setText(str(day1_str))
        self.day2_value.setText(str(day2_str))
        self.path_group_folder_value.setText(os.path.abspath(os.getcwd()) + '/' + self.folder_name())

    def folder_name(self):
        second_day = self.day2_value.text()
        day = str(second_day[0:2])
        month = str(second_day[3:5])
        year = str(second_day[6:10])
        folder_name = day + month + year + ' FMD'
        return folder_name

    def select_group_folder(self):
        directory = QtWidgets.QFileDialog.getExistingDirectory(self, "Select folder")
        if directory:
            directory = directory + '/' + self.folder_name()
            if not os.path.isdir(directory):
                self.path_group_folder_value.setText('Path: ' + directory)

    def generate_group_folder(self):
        self.status_label.setText('')
        first_date = self.day1_value.text()
        second_date = self.day2_value.text()
        names = self.get_students_names()
        try:
            folder_path = self.path_group_folder_value.text()
            if not os.path.isdir(folder_path):
                os.makedirs(folder_path)
            self.create_attendance_doc(first_date, second_date, names, folder_path)
            for name in names:
                self.generate_term_doc(folder_path, second_date, name)
            self.status_label.setText(self.status_label.text() + ' Term files created.')
            for name in names:
                self.generate_score_files(name, folder_path + '/', first_date, second_date)
            self.status_label.setText(self.status_label.text() + ' Score files created.')
            data = [first_date, second_date, folder_path + '/', names]
            self.genetate_adf06(data)
            self.status_label_2.setText(self.status_label_2.text() + ' ADF06 files created.')
            self.genetate_adf21(data)
            self.status_label_2.setText(self.status_label_2.text() + ' ADF21 files created.')
            self.genetate_adf03(data)
            self.status_label_2.setText(self.status_label_2.text() + ' ADF03 files created.')
            self.status_label_2.setText(self.status_label_2.text() + ' Job done.')
        except:
            self.status_label_2.setText(self.status_label_2.text() + 'Error.')

    def create_attendance_doc(self, first_date, second_date, names, folder_path):
        attendance_file = self.work_file_path + 'attendance.docx'
        attendance = Document(attendance_file)
        attendance_table = attendance.tables[1]
        j = 1
        for name in names:
            attendance_table.cell(j, 1).text = name
            attendance_table.cell(j, 2).text = first_date
            run = attendance_table.cell(j, 2).paragraphs[0].runs[0]
            font = run.font
            font.size = shared.Pt(8)
            attendance_table.cell(j, 3).text = second_date
            run2 = attendance_table.cell(j, 3).paragraphs[0].runs[0]
            font2 = run2.font
            font2.size = shared.Pt(8)
            j = j + 1

        attendance_table2 = attendance.tables[2]
        attendance_table2.cell(0, 1).text = second_date
        attendance.save(folder_path + '/attendance.docx')
        self.status_label.setText(self.status_label.text() + ' Attendance file created.')

    def get_students_names(self):
        names = []
        try:
            names_doc = Document(self.name_file_path)
            names_doc_table = names_doc.tables[0]
            rows = len(names_doc_table.rows)
            i = 1
            while i != rows:
                if names_doc_table.cell(i, 0).text:
                    temp = names_doc_table.cell(i, 0).text + ' ' + names_doc_table.cell(i, 1).text
                    names.append(temp)
                i = i + 1
            self.status_label.setText(self.status_label.text() + 'Get student names.')
            return names
        except:
            self.status_label.setText(self.status_label.text() + 'Problem with student names.')
            return names

    def generate_score_files(self, student_name, path, first_date, second_date):
        day = int(second_date[0:2])
        month = int(second_date[3:5])
        year = int(second_date[6:10])
        dt = datetime.date(year, month, day)
        wk = dt.isocalendar()[1]
        course_number = str(year) + '-W' + str(wk)
        date = datetime.datetime(year, month, day, 00, 00)
        wb = openpyxl.load_workbook(self.work_file_path + 'score.xlsx')
        sheet = wb['Ark1']
        sheet['D3'] = course_number
        sheet['D4'] = date
        sheet['D5'] = student_name
        sheet['D6'] = 'OMTC'
        sheet['D15'] = 'V.V.'
        file_name = self.generate_score_file_name(student_name, first_date, second_date)
        wb.save(path + file_name)

    def generate_score_file_name(self, student_name, first_date, second_date):
        day1 = int(first_date[0:2])
        month1 = int(first_date[3:5])
        year1 = int(first_date[6:10])
        day2 = int(second_date[0:2])
        month2 = int(second_date[3:5])
        year2 = int(second_date[6:10])
        dt = datetime.date(year2, month2, day2)
        wk = dt.isocalendar()[1]
        if month1 < 10:
            month1 = '0' + str(month1)
        if day1 < 10:
            day1 = '0' + str(day1)
        if month2 < 10:
            month2 = '0' + str(month2)
        if day2 < 10:
            day2 = '0' + str(day2)
        name = 'OMTC_W' + str(wk) + '_' + str(year1) + '-' + str(month1) + '-' + \
               str(day1) + '_' + str(year2) + '-' + str(month2) + '-' + str(day2) + '_SCORE_' + student_name + '.xlsx'
        return name

    def generate_test_file_name(self, student_name, first_date, second_date):
        day1 = int(first_date[0:2])
        month1 = int(first_date[3:5])
        year1 = int(first_date[6:10])
        day2 = int(second_date[0:2])
        month2 = int(second_date[3:5])
        year2 = int(second_date[6:10])
        dt = datetime.date(year2, month2, day2)
        wk = dt.isocalendar()[1]
        if month1 < 10:
            month1 = '0' + str(month1)
        if day1 < 10:
            day1 = '0' + str(day1)
        if month2 < 10:
            month2 = '0' + str(month2)
        if day2 < 10:
            day2 = '0' + str(day2)
        name = 'OMTC_W' + str(wk) + '_' + str(year1) + '-' + str(month1) + '-' + \
               str(day1) + '_' + str(year2) + '-' + str(month2) + '-' + str(day2) + '_TEST_' + student_name + '.pdf'
        return name

    def generate_term_doc(self, folder_path, day2, student_name):
        term = self.work_file_path + 'term.docx'
        term_doc = Document(term)
        res = term_doc.paragraphs[10].runs[0].font.name
        term_doc.paragraphs[10].text = 'Date: ' + day2
        term_doc.paragraphs[12].text = 'Name in capital letters: ' + student_name
        term_doc.paragraphs[10].runs[0].font.name = res
        term_doc.paragraphs[12].runs[0].font.name = res
        term_doc.save(folder_path + '/term_' + student_name + '.docx')

    def genetate_adf06(self, data):
        date_start = data[0]
        date_end = data[1]
        path = data[2]
        names = data[3]
        file_name = 'adf06.pdf'
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        can.drawString(640, 527, date_start)
        can.drawString(740, 527, date_end)
        can.setFontSize(10)
        i = 60
        j = 203
        for name in names:
            can.drawString(i, j, name)
            j = j - 13
        can.save()
        packet.seek(0)
        new_pdf = PdfFileReader(packet)
        existing_pdf = PdfFileReader(open(self.work_file_path + "adf06.pdf", "rb"))
        output = PdfFileWriter()
        page = existing_pdf.getPage(0)
        page.mergePage(new_pdf.getPage(0))
        output.addPage(page)
        outputStream = open(path + file_name, "wb")
        output.write(outputStream)
        outputStream.close()

    def genetate_adf21(self, data):
        date_start = data[0]
        date_end = data[1]
        path = data[2]
        file_name = 'adf21.pdf'
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        can.setFontSize(10)
        can.drawString(665, 510, date_start)
        can.drawString(750, 510, date_end)
        can.rotate(90)
        can.drawString(430, -515, date_start)
        can.drawString(430, -537, date_end)
        can.save()
        packet.seek(0)
        new_pdf = PdfFileReader(packet)
        existing_pdf = PdfFileReader(open(self.work_file_path + "adf21.pdf", "rb"))
        output = PdfFileWriter()
        page = existing_pdf.getPage(0)
        page.mergePage(new_pdf.getPage(0))
        output.addPage(page)
        outputStream = open(path + file_name, "wb")
        output.write(outputStream)
        outputStream.close()

    def genetate_adf03(self, data):
        date_end = data[1]
        path = data[2]
        file_name = 'adf03.pdf'
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        can.drawString(95, 145, date_end)
        can.save()
        packet.seek(0)
        new_pdf = PdfFileReader(packet)
        existing_pdf = PdfFileReader(open(self.work_file_path + "adf03.pdf", "rb"))
        output = PdfFileWriter()
        page = existing_pdf.getPage(0)
        page.mergePage(new_pdf.getPage(0))
        output.addPage(page)
        outputStream = open(path + file_name, "wb")
        output.write(outputStream)
        outputStream.close()

    def generate_test_files(self):
        self.status_label.setText('')
        self.status_label_2.setText('')
        first_date = self.day1_value.text()
        second_date = self.day2_value.text()
        names = self.get_students_names()
        folder_path = self.path_group_folder_value.text()
        print(folder_path)
        count = 1
        for name in names:
            file_name = self.generate_test_file_name(name, first_date, second_date)
            output = PdfFileWriter()
            front_page_file = PdfFileReader(open(folder_path + '/' + str(count) + ".pdf", "rb"))
            test_file = PdfFileReader(open(self.work_file_path + "test " + str(count) + ".pdf", "rb"))
            output.addPage(front_page_file.getPage(0))
            for x in range(5):
                output.addPage(test_file.getPage(x))
            outputStream = open(folder_path + '/' + file_name, "wb")
            output.write(outputStream)
            outputStream.close()

            count = count + 1
        self.status_label_2.setText(self.status_label_2.text() + ' Job done.')


def main():
    app = QtWidgets.QApplication(sys.argv)
    window = MainApp()
    window.show()
    app.exec_()


if __name__ == '__main__':
    main()
